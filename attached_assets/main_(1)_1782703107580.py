"""
EduBharat — single-file Replit app

MODES (deliberately kept separate, per your request):
  1. ENGLISH GURU  — live conversation practice (chat + TTS). Routes: /api/guru/*, /api/tts
  2. JOBS          — personalized job search.                  Routes: /api/jobs/*
  3. NEWS FEED     — personalized news feed.                   Routes: /api/news/*
  4. RESUME TAILOR — upload, tailor to a target job, download.  Routes: /api/resume/*
  5. LEARNING JOURNEY — research-based lesson sequencing.       Routes: /api/journey/*

SETUP ON REPLIT:
  1. Create a Python Repl, paste this whole file in as main.py
  2. In the Shell tab, run:
       pip install flask edge-tts anthropic pdfplumber python-docx
  3. In Replit's "Secrets" tab, add: ANTHROPIC_API_KEY = <your key>
     (only needed for English Guru chat + Resume Tailor — TTS, Jobs, News,
     and Learning Journey all work without it)
  4. Click Run, open the webview

WHAT'S REAL VS. PLACEHOLDER (read this before you judge "is it working"):
  - TTS, the SM-2 spaced-repetition math, the weighted job/news scoring, and
    the full resume parse->tailor->render->download pipeline are genuinely
    working code, not stubs.
  - MOCK_JOBS and MOCK_NEWS are small sample datasets so you can see scoring
    in action. Swap them for your real Rozgar Samachar job feed / news
    source — the scoring functions (score_job, score_news) don't need to
    change, they just need real data and real profile fields to read from.
  - User profile and learning progress are stored in memory (PROGRESS dict)
    and reset every time the Repl restarts. Swap get_user_progress() for a
    real database call once this is wired into your actual login system.
"""

import os
import io
import json
import uuid
import asyncio
from datetime import date, timedelta

from flask import Flask, request, jsonify, send_file, Response

import edge_tts
import anthropic
import pdfplumber
from docx import Document

app = Flask(__name__)
client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from Replit Secrets


def _clean_json_text(text):
    """Strip ```json fences if the model adds them despite instructions not to."""
    text = text.strip()
    if text.startswith("```"):
        parts = text.split("```")
        text = parts[1] if len(parts) >= 2 else text
        if text.startswith("json"):
            text = text[4:]
    return text.strip()


# =====================================================================
# MODE 1: ENGLISH GURU — live conversation practice
# Kept on its own routes (/api/guru/*) and its own panel in the UI,
# separate from the career tools below.
# =====================================================================

VOICES = {
    "en-in":   "en-IN-PrabhatNeural",
    "en-in-f": "en-IN-NeerjaNeural",
    "hi":      "hi-IN-MadhurNeural",
    "hi-f":    "hi-IN-SwaraNeural",
    "ta":      "ta-IN-ValluvarNeural",
    "ta-f":    "ta-IN-PallaviNeural",
    "te":      "te-IN-MohanNeural",
    "te-f":    "te-IN-ShrutiNeural",
    "bn":      "bn-IN-BashkarNeural",
    "bn-f":    "bn-IN-TanishaaNeural",
    "mr":      "mr-IN-ManoharNeural",
    "mr-f":    "mr-IN-AarohiNeural",
    "gu":      "gu-IN-NiranjanNeural",
    "gu-f":    "gu-IN-DhwaniNeural",
    "kn":      "kn-IN-GaganNeural",
    "kn-f":    "kn-IN-SapnaNeural",
    "ml":      "ml-IN-MidhunNeural",
    "ml-f":    "ml-IN-SobhanaNeural",
    "ur":      "ur-IN-SalmanNeural",
    "ur-f":    "ur-IN-GulNeural",
    "pa":      "pa-IN-OjasNeural",
    "pa-f":    "pa-IN-VaaniNeural",
    "or":      "or-IN-SukantNeural",
    "or-f":    "or-IN-SubhasiniNeural",
    "as":      "as-IN-PriyomNeural",
    "as-f":    "as-IN-YashicaNeural",
}

OUTPUT_DIR = "tts_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

GURU_SYSTEM_PROMPT = """You are English Guru, a friendly, patient spoken-English tutor for
Indian learners. Keep replies short (1-3 sentences) since they'll be spoken aloud.
Gently correct grammar mistakes by modeling the correct form in your reply, without
lecturing about the mistake. Ask a natural follow-up question to keep the conversation
going, like a real tutor would."""


@app.route("/api/voices", methods=["GET"])
def list_voices():
    return jsonify(list(VOICES.keys()))


@app.route("/api/tts", methods=["POST"])
def tts():
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()
    lang = data.get("lang", "en-in")
    if not text:
        return jsonify({"error": "Missing 'text' field"}), 400

    voice = VOICES.get(lang, VOICES["en-in"])
    filepath = os.path.join(OUTPUT_DIR, f"{uuid.uuid4().hex}.mp3")

    async def generate():
        await edge_tts.Communicate(text, voice).save(filepath)

    try:
        asyncio.run(generate())
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    response = send_file(filepath, mimetype="audio/mpeg")

    @response.call_on_close
    def cleanup():
        if os.path.exists(filepath):
            os.remove(filepath)

    return response


@app.route("/api/guru/chat", methods=["POST"])
def guru_chat():
    data = request.get_json(force=True) or {}
    history = data.get("history", [])
    if not history:
        return jsonify({"error": "Missing 'history'"}), 400

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=200,
            system=GURU_SYSTEM_PROMPT,
            messages=history,
        )
        reply = response.content[0].text
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    return jsonify({"reply": reply})


# =====================================================================
# MODE 2: JOBS — personalized search (hard filters happen client-side
# via the form; this scores+ranks what matches)
# =====================================================================

MOCK_JOBS = [
    {"id": "j1", "title": "Data Entry Operator", "company": "Sahara Infotech",
     "location": "Kanpur", "remote": False, "sector": "government",
     "required_skills": ["typing", "ms office", "communication"],
     "min_experience": 0, "max_experience": 2, "days_since_posted": 1},
    {"id": "j2", "title": "Bank PO", "company": "State Bank Recruitment",
     "location": "Lucknow", "remote": False, "sector": "banking",
     "required_skills": ["banking", "communication", "reasoning"],
     "min_experience": 0, "max_experience": 3, "days_since_posted": 3},
    {"id": "j3", "title": "Junior Software Developer", "company": "TechNova",
     "location": "Bengaluru", "remote": True, "sector": "technology",
     "required_skills": ["python", "sql", "git"],
     "min_experience": 0, "max_experience": 2, "days_since_posted": 0},
    {"id": "j4", "title": "Customer Support Executive", "company": "Vodafone Idea",
     "location": "Kanpur", "remote": False, "sector": "telecom",
     "required_skills": ["communication", "english", "crm"],
     "min_experience": 0, "max_experience": 4, "days_since_posted": 2},
    {"id": "j5", "title": "Sales Associate", "company": "Reliance Retail",
     "location": "Kanpur", "remote": False, "sector": "retail",
     "required_skills": ["sales", "communication", "customer service"],
     "min_experience": 0, "max_experience": 3, "days_since_posted": 5},
]


def score_job(profile, job):
    """Weighted match score. Tune the weights once you see real click data."""
    score = 0.0
    user_skills = {s.lower() for s in profile.get("skills", [])}
    job_skills = {s.lower() for s in job["required_skills"]}
    if job_skills:
        overlap = len(user_skills & job_skills) / len(job_skills)
        score += overlap * 50  # skill match dominates the ranking

    years = profile.get("years_experience", 0)
    if job["min_experience"] <= years <= job["max_experience"] + 2:
        score += 20

    preferred_location = (profile.get("preferred_location") or "").lower()
    if preferred_location and job["location"].lower() == preferred_location:
        score += 15
    elif job["remote"]:
        score += 10

    if job["sector"] in profile.get("interested_sectors", []):
        score += 10

    score += max(0, 5 - job["days_since_posted"])  # small recency boost
    return score


@app.route("/api/jobs/search", methods=["POST"])
def jobs_search():
    profile = request.get_json(force=True) or {}
    scored = [{**job, "score": score_job(profile, job)} for job in MOCK_JOBS]
    scored.sort(key=lambda j: j["score"], reverse=True)
    return jsonify({"results": scored})


# =====================================================================
# MODE 3: NEWS FEED — personalized
# =====================================================================

MOCK_NEWS = [
    {"id": "n1", "title": "SBI PO 2026 notification released", "sector": "banking", "days_old": 0},
    {"id": "n2", "title": "UP Police constable recruitment exam date out", "sector": "government", "days_old": 1},
    {"id": "n3", "title": "IT hiring rebounds in Q2 2026", "sector": "technology", "days_old": 2},
    {"id": "n4", "title": "Retail sector adds 50,000 jobs this festive season", "sector": "retail", "days_old": 4},
    {"id": "n5", "title": "Telecom companies expand rural hiring drives", "sector": "telecom", "days_old": 1},
]


def score_news(profile, article):
    score = 0.0
    if article["sector"] in profile.get("interested_sectors", []):
        score += 40
    score += max(0, 10 - article["days_old"])  # recency decay
    return score


@app.route("/api/news/feed", methods=["POST"])
def news_feed():
    profile = request.get_json(force=True) or {}
    scored = [{**a, "score": score_news(profile, a)} for a in MOCK_NEWS]
    scored.sort(key=lambda a: a["score"], reverse=True)
    return jsonify({"results": scored})


# =====================================================================
# MODE 4: RESUME TAILOR
# Pipeline: extract text -> structure to JSON -> tailor to target job
# -> render to .docx -> serve as a real downloadable attachment
# =====================================================================

def extract_text(file_stream, filename):
    if filename.lower().endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text
    elif filename.lower().endswith(".docx"):
        doc = Document(file_stream)
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Unsupported file type — upload a .pdf or .docx")


def parse_resume_to_json(raw_text):
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[{
            "role": "user",
            "content": (
                "Extract this resume into JSON with keys: name, contact, summary, "
                "skills (list), experience (list of {title, company, dates, bullets}), "
                "education (list). Return ONLY valid JSON, no preamble, no markdown fences.\n\n"
                f"RESUME TEXT:\n{raw_text}"
            ),
        }],
    )
    return json.loads(_clean_json_text(response.content[0].text))


def tailor_resume_json(resume_json, target_job, experience_level):
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": (
                "Tailor this resume JSON for the target job below.\n"
                "Rules:\n"
                "- Reorder 'skills' so ones matching the target job come first\n"
                "- Rewrite 'summary' to highlight fit for this role and experience level\n"
                "- Reword bullets to emphasize relevant achievements (don't invent facts)\n"
                f"- Experience level selected: {experience_level} (freshers -> emphasize "
                "education/projects; senior -> emphasize leadership/impact)\n\n"
                f"TARGET JOB: {target_job}\n\n"
                f"RESUME JSON:\n{json.dumps(resume_json)}\n\n"
                "Return ONLY the tailored JSON, same schema, no preamble, no markdown fences."
            ),
        }],
    )
    return json.loads(_clean_json_text(response.content[0].text))


def render_docx(data):
    doc = Document()
    doc.add_heading(data.get("name", "Resume"), level=0)
    if data.get("contact"):
        doc.add_paragraph(data["contact"])
    if data.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data["summary"])
    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data["skills"]))
    if data.get("experience"):
        doc.add_heading("Experience", level=1)
        for job in data["experience"]:
            doc.add_paragraph(
                f"{job.get('title', '')} — {job.get('company', '')} ({job.get('dates', '')})",
                style="Heading 3",
            )
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")
    if data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            doc.add_paragraph(edu)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/api/resume/tailor", methods=["POST"])
def resume_tailor():
    if "resume" not in request.files:
        return jsonify({"error": "Missing 'resume' file"}), 400
    file = request.files["resume"]
    target_job = request.form.get("target_job", "").strip()
    experience_level = request.form.get("experience_level", "fresher")

    if not target_job:
        return jsonify({"error": "Missing 'target_job'"}), 400

    try:
        raw_text = extract_text(file.stream, file.filename)
        structured = parse_resume_to_json(raw_text)
        tailored = tailor_resume_json(structured, target_job, experience_level)
        docx_buffer = render_docx(tailored)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    return send_file(
        docx_buffer,
        as_attachment=True,                      # <- this is what makes it actually download
        download_name="tailored_resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# =====================================================================
# MODE 5: LEARNING JOURNEY — research-based sequencing
#   - Spaced repetition (SM-2 algorithm, the same family of algorithm
#     Anki uses, grounded in the Ebbinghaus forgetting-curve research)
#     decides WHEN a lesson should come back for review.
#   - Interleaving (mixing skill types instead of blocking one type at
#     a time — supported by Rohrer & Taylor's interleaving research)
#     decides the ORDER lessons are presented in.
#   - Retrieval practice (a scored attempt, not passive re-reading —
#     per Roediger & Karpicke's testing-effect research) is what drives
#     the spaced-repetition update below.
# =====================================================================

LESSON_BANK = [
    {"id": "l1", "title": "Greetings & introductions", "skill_type": "vocabulary"},
    {"id": "l2", "title": "Present simple tense", "skill_type": "grammar"},
    {"id": "l3", "title": "Listening: ordering at a shop", "skill_type": "listening"},
    {"id": "l4", "title": "Speaking: describe your day", "skill_type": "speaking"},
    {"id": "l5", "title": "Workplace vocabulary", "skill_type": "vocabulary"},
    {"id": "l6", "title": "Past simple tense", "skill_type": "grammar"},
    {"id": "l7", "title": "Listening: a job interview", "skill_type": "listening"},
    {"id": "l8", "title": "Speaking: answer interview questions", "skill_type": "speaking"},
]

# In-memory per-learner spaced-repetition state.
# Swap for a real DB table: (user_id, lesson_id) -> {ease, interval, repetitions, due_date, last_score}
PROGRESS = {}
DEMO_USER = "demo_user"  # replace with the logged-in user's real id


def sm2_update(ease, interval, repetitions, quality):
    """SM-2 spaced repetition. quality is 0-5 (0 = forgot completely, 5 = perfect recall)."""
    if quality < 3:
        repetitions = 0
        interval = 1
    else:
        if repetitions == 0:
            interval = 1
        elif repetitions == 1:
            interval = 6
        else:
            interval = round(interval * ease)
        repetitions += 1

    ease = ease + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02))
    ease = max(1.3, ease)
    return ease, interval, repetitions


def get_user_progress(user_id):
    return PROGRESS.setdefault(user_id, {})


@app.route("/api/journey/next", methods=["GET"])
def journey_next():
    progress = get_user_progress(DEMO_USER)
    today = date.today()
    due, new = [], []

    for lesson in LESSON_BANK:
        state = progress.get(lesson["id"])
        if state is None:
            new.append(lesson)
        elif date.fromisoformat(state["due_date"]) <= today:
            due.append(lesson)

    # Due reviews first (spaced repetition), then new material — then
    # interleave across skill types so the same skill never repeats back-to-back.
    candidates = due + new
    by_type = {}
    for lesson in candidates:
        by_type.setdefault(lesson["skill_type"], []).append(lesson)

    interleaved = []
    while any(by_type.values()):
        for skill_type in list(by_type.keys()):
            bucket = by_type[skill_type]
            if bucket:
                interleaved.append(bucket.pop(0))
            if not bucket:
                del by_type[skill_type]

    result = []
    for lesson in interleaved[:5]:
        state = progress.get(lesson["id"])
        status = "due for review" if state else "new lesson"
        result.append({**lesson, "status": status})

    return jsonify({"lessons": result})


@app.route("/api/journey/submit-result", methods=["POST"])
def journey_submit_result():
    data = request.get_json(force=True) or {}
    lesson_id = data.get("lesson_id")
    score = float(data.get("score", 0))  # 0-100, e.g. from a quiz

    if lesson_id not in {l["id"] for l in LESSON_BANK}:
        return jsonify({"error": "Unknown lesson_id"}), 400

    progress = get_user_progress(DEMO_USER)
    state = progress.get(lesson_id, {"ease": 2.5, "interval": 0, "repetitions": 0})

    quality = round(score / 100 * 5)
    ease, interval, repetitions = sm2_update(
        state["ease"], state["interval"], state["repetitions"], quality
    )
    due_date = date.today() + timedelta(days=interval)

    progress[lesson_id] = {
        "ease": ease,
        "interval": interval,
        "repetitions": repetitions,
        "due_date": due_date.isoformat(),
        "last_score": score,
    }

    return jsonify({"ok": True, "next_review": due_date.isoformat()})


# =====================================================================
# HOME PAGE — single test UI. English Guru gets its own visually distinct,
# always-first panel; Jobs / News / Resume / Journey are grouped separately
# as "career tools" tabs.
# =====================================================================

PAGE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<title>EduBharat</title>
<style>
  body { font-family: system-ui, sans-serif; max-width: 760px; margin: 30px auto; padding: 0 16px; color: #1a1a1a; }
  h1 { margin-bottom: 4px; }
  .tabs { display: flex; gap: 8px; margin: 20px 0; flex-wrap: wrap; }
  .tab-btn { padding: 10px 16px; border: 1px solid #ccc; background: #f5f5f5; border-radius: 8px; cursor: pointer; font-size: 14px; }
  .tab-btn.active { background: #2d6cdf; color: white; border-color: #2d6cdf; }
  .tab-btn.guru.active { background: #16a37a; border-color: #16a37a; }
  .panel { display: none; border: 1px solid #e2e2e2; border-radius: 10px; padding: 20px; }
  .panel.active { display: block; }
  .guru-panel { background: #f3fbf8; border-color: #16a37a; }
  textarea, input, select { width: 100%; font-size: 15px; padding: 8px; margin: 6px 0; box-sizing: border-box; }
  button { padding: 8px 14px; font-size: 15px; cursor: pointer; border-radius: 6px; border: 1px solid #2d6cdf; background: #2d6cdf; color: white; margin-top: 6px; }
  button.secondary { background: white; color: #2d6cdf; }
  .chat-log { max-height: 260px; overflow-y: auto; border: 1px solid #ddd; border-radius: 8px; padding: 10px; margin: 10px 0; background: white; }
  .msg-user { text-align: right; color: #2d6cdf; margin: 6px 0; }
  .msg-bot { text-align: left; color: #16a37a; margin: 6px 0; }
  .job-card, .news-card, .lesson-card { border: 1px solid #e2e2e2; border-radius: 8px; padding: 10px; margin: 8px 0; }
  .score-tag { font-size: 12px; background: #eef; padding: 2px 6px; border-radius: 4px; }
  small.note { color: #888; }
</style>
</head>
<body>

<h1>EduBharat</h1>
<small class="note">English Guru is a separate live-conversation mode from the career tools below.</small>

<div class="tabs">
  <div class="tab-btn guru active" onclick="showTab('guru', event)">English Guru</div>
  <div class="tab-btn" onclick="showTab('jobs', event)">Jobs</div>
  <div class="tab-btn" onclick="showTab('news', event)">News Feed</div>
  <div class="tab-btn" onclick="showTab('resume', event)">Resume Tailor</div>
  <div class="tab-btn" onclick="showTab('journey', event)">Learning Journey</div>
</div>

<div id="panel-guru" class="panel guru-panel active">
  <h3>English Guru — Live Conversation Practice</h3>
  <select id="guruVoice">{voice_options}</select>
  <div id="chatLog" class="chat-log"></div>
  <textarea id="guruInput" placeholder="Type or use the mic..."></textarea>
  <button onclick="micInput()" class="secondary">Speak (mic)</button>
  <button onclick="sendGuru()">Send</button>
</div>

<div id="panel-jobs" class="panel">
  <h3>Personalized Job Search</h3>
  <input id="jobSkills" placeholder="Your skills, comma separated e.g. excel, sales, communication"/>
  <input id="jobLocation" placeholder="Preferred location e.g. Kanpur"/>
  <input id="jobExp" type="number" placeholder="Years of experience" value="0"/>
  <button onclick="searchJobs()">Search Jobs</button>
  <div id="jobResults"></div>
</div>

<div id="panel-news" class="panel">
  <h3>Personalized News Feed</h3>
  <input id="newsSectors" placeholder="Sectors of interest e.g. government, banking, technology"/>
  <button onclick="loadNews()">Load Feed</button>
  <div id="newsResults"></div>
</div>

<div id="panel-resume" class="panel">
  <h3>Resume Tailor</h3>
  <input type="file" id="resumeFile" accept=".pdf,.docx"/>
  <input id="targetJob" placeholder="Target job title / description"/>
  <select id="expLevel">
    <option value="fresher">Fresher</option>
    <option value="mid">Mid-level</option>
    <option value="senior">Senior</option>
  </select>
  <button onclick="tailorResume()">Generate Tailored Resume</button>
  <div id="resumeStatus"></div>
</div>

<div id="panel-journey" class="panel">
  <h3>My Learning Journey</h3>
  <small class="note">Lesson order uses spaced repetition + interleaving, not a fixed playlist.</small>
  <button onclick="loadJourney()">Get My Next Lessons</button>
  <div id="journeyResults"></div>
</div>

<script>
function showTab(name, evt) {
  document.querySelectorAll('.panel').forEach(p => p.classList.remove('active'));
  document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
  document.getElementById('panel-' + name).classList.add('active');
  evt.currentTarget.classList.add('active');
}

let chatHistory = [];

function micInput() {
  const Rec = window.SpeechRecognition || window.webkitSpeechRecognition;
  if (!Rec) { alert('Speech recognition not supported in this browser — try Chrome.'); return; }
  const rec = new Rec();
  rec.lang = 'en-IN';
  rec.onresult = (e) => { document.getElementById('guruInput').value = e.results[0][0].transcript; };
  rec.start();
}

async function sendGuru() {
  const input = document.getElementById('guruInput');
  const text = input.value.trim();
  if (!text) return;
  const log = document.getElementById('chatLog');
  log.innerHTML += `<div class="msg-user">${text}</div>`;
  chatHistory.push({ role: 'user', content: text });
  input.value = '';

  const res = await fetch('/api/guru/chat', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ history: chatHistory })
  });
  const data = await res.json();
  const reply = data.reply || ('Error: ' + (data.error || 'unknown'));
  chatHistory.push({ role: 'assistant', content: reply });
  log.innerHTML += `<div class="msg-bot">${reply}</div>`;
  log.scrollTop = log.scrollHeight;

  if (data.reply) {
    const voice = document.getElementById('guruVoice').value;
    const ttsRes = await fetch('/api/tts', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ text: reply, lang: voice })
    });
    if (ttsRes.ok) {
      const blob = await ttsRes.blob();
      new Audio(URL.createObjectURL(blob)).play();
    }
  }
}

async function searchJobs() {
  const skills = document.getElementById('jobSkills').value.split(',').map(s => s.trim()).filter(Boolean);
  const location = document.getElementById('jobLocation').value.trim();
  const exp = parseFloat(document.getElementById('jobExp').value) || 0;

  const res = await fetch('/api/jobs/search', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ skills, preferred_location: location, years_experience: exp })
  });
  const data = await res.json();
  document.getElementById('jobResults').innerHTML = data.results.map(j =>
    `<div class="job-card"><b>${j.title}</b> — ${j.company} (${j.location})<br/>
     <span class="score-tag">match score: ${j.score.toFixed(1)}</span></div>`
  ).join('') || '<p>No matches found.</p>';
}

async function loadNews() {
  const sectors = document.getElementById('newsSectors').value.split(',').map(s => s.trim()).filter(Boolean);
  const res = await fetch('/api/news/feed', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ interested_sectors: sectors })
  });
  const data = await res.json();
  document.getElementById('newsResults').innerHTML = data.results.map(n =>
    `<div class="news-card"><b>${n.title}</b><br/><small>${n.sector} · ${n.days_old}d ago</small></div>`
  ).join('') || '<p>No articles found.</p>';
}

async function tailorResume() {
  const fileInput = document.getElementById('resumeFile');
  const targetJob = document.getElementById('targetJob').value.trim();
  const expLevel = document.getElementById('expLevel').value;
  const status = document.getElementById('resumeStatus');

  if (!fileInput.files.length || !targetJob) {
    status.textContent = 'Please upload a resume and enter a target job.';
    return;
  }

  const formData = new FormData();
  formData.append('resume', fileInput.files[0]);
  formData.append('target_job', targetJob);
  formData.append('experience_level', expLevel);

  status.textContent = 'Tailoring your resume...';
  try {
    const res = await fetch('/api/resume/tailor', { method: 'POST', body: formData });
    if (!res.ok) {
      const err = await res.json();
      throw new Error(err.error || 'Failed');
    }
    const blob = await res.blob();
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;
    a.download = 'tailored_resume.docx';
    document.body.appendChild(a);
    a.click();
    a.remove();
    status.textContent = 'Done — check your downloads.';
  } catch (e) {
    status.textContent = 'Error: ' + e.message;
  }
}

async function loadJourney() {
  const res = await fetch('/api/journey/next');
  const data = await res.json();
  document.getElementById('journeyResults').innerHTML = data.lessons.map(l =>
    `<div class="lesson-card">
       <b>${l.title}</b> <span class="score-tag">${l.skill_type}</span><br/>
       <small>${l.status}</small><br/>
       Score: <input type="number" min="0" max="100" id="score-${l.id}" style="width:70px;display:inline-block;"/>
       <button class="secondary" onclick="submitResult('${l.id}')">Mark complete</button>
     </div>`
  ).join('') || '<p>No lessons due right now — nice work staying on top of reviews.</p>';
}

async function submitResult(lessonId) {
  const score = parseFloat(document.getElementById('score-' + lessonId).value) || 0;
  await fetch('/api/journey/submit-result', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ lesson_id: lessonId, score })
  });
  loadJourney();
}
</script>
</body>
</html>
"""


@app.route("/")
def home():
    voice_options = "\n".join(f'<option value="{k}">{k}</option>' for k in VOICES)
    return Response(PAGE.replace("{voice_options}", voice_options), mimetype="text/html")


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
